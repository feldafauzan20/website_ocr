"""
Bot Telegram untuk mengubah gambar tabel menjadi file JSON menggunakan Gemini Vision API.
"""
import asyncio
import json
import logging
import os
import uuid
from datetime import datetime
import docx
import pdfplumber
from dotenv import load_dotenv
from telegram import Update
from telegram.constants import ParseMode
from telegram.ext import (Application, CommandHandler, ContextTypes,
                          MessageHandler, filters)

import gemini_vision_extractor

load_dotenv()

logging.basicConfig(
    format="%(asctime)s - %(name)s - [%(levelname)s] - %(message)s", level=logging.INFO
)
logger = logging.getLogger(__name__)

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not TELEGRAM_BOT_TOKEN:
    raise ValueError("TELEGRAM_BOT_TOKEN tidak ditemukan di file .env")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY tidak ditemukan di file .env")


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    await update.message.reply_html(
        rf"Halo {user.mention_html()}! Kirimkan gambar tabel. Saya akan mengonversinya menjadi file JSON menggunakan AI.",
    )

def pdf_to_json(pdf_path):
    """
    Ekstrak tabel dari PDF dan konversi ke JSON array of objects.
    Hanya mengambil tabel pertama di halaman pertama.
    """
    logger.info(f"Membuka PDF: {pdf_path}")
    with pdfplumber.open(pdf_path) as pdf:
        first_page = pdf.pages[0]
        tables = first_page.extract_tables()
        if not tables:
            logger.warning(f"Tidak ditemukan tabel di PDF: {pdf_path}")
            return []
        table = tables[0]
        headers = table[0]
        data = []
        for row in table[1:]:
            obj = {}
            for i, cell in enumerate(row):
                key = headers[i] if i < len(headers) else f"col_{i+1}"
                obj[key] = cell if cell not in [None, ""] else None
            data.append(obj)
        logger.info(f"Berhasil mengekstrak {len(data)} baris dari PDF.")
        return data

def docx_to_json(docx_path):
    """
    Ekstrak tabel dari DOCX dan konversi ke JSON array of objects.
    Hanya mengambil tabel pertama.
    """
    logger.info(f"Membuka DOCX: {docx_path}")
    doc = docx.Document(docx_path)
    if not doc.tables:
        logger.warning(f"Tidak ditemukan tabel di DOCX: {docx_path}")
        return []
    table = doc.tables[0]
    rows = list(table.rows)
    headers = [cell.text.strip() for cell in rows[0].cells]
    data = []
    for row in rows[1:]:
        obj = {}
        for i, cell in enumerate(row.cells):
            key = headers[i] if i < len(headers) else f"col_{i+1}"
            value = cell.text.strip()
            obj[key] = value if value else None
        data.append(obj)
    logger.info(f"Berhasil mengekstrak {len(data)} baris dari DOCX.")
    return data

async def process_pdf_and_send_json(context, chat_id, temp_pdf_path, message_id, original_base_filename):
    output_json_path = None
    try:
        await context.bot.edit_message_text(
            text="⏳ Memproses PDF untuk menghasilkan JSON...",
            chat_id=chat_id,
            message_id=message_id
        )
        data = pdf_to_json(temp_pdf_path)
        data = fix_empty_key(data, new_key="Akun")
        if not data:
            await context.bot.edit_message_text(
                text="⚠️ Tidak ditemukan tabel pada PDF.",
                chat_id=chat_id,
                message_id=message_id
            )
            logger.info("Tidak ada data tabel yang diekstrak dari PDF.")
            return
        
        # Gunakan nama file asli sebagai nama file JSON
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        unique_id = uuid.uuid4().hex[:8] 
        output_json_path = os.path.join("output", f"{original_base_filename}_{timestamp}_{unique_id}.json")
        logger.info(f"Akan menyimpan JSON ke: {output_json_path}")
        
        with open(output_json_path, "w", encoding="utf-8") as f:
            json_to_write = json.dumps(data, ensure_ascii=False, indent=2)
            f.write(json_to_write)
        logger.info(f"JSON berhasil ditulis ke: {output_json_path}")

        await context.bot.edit_message_text(
            text="✅ JSON berhasil dibuat. Mengirim file ke Anda...",
            chat_id=chat_id,
            message_id=message_id
        )
        await context.bot.send_document(
            chat_id=chat_id,
            document=open(output_json_path, 'rb'),
            filename=os.path.basename(output_json_path), # Pastikan nama file benar untuk Telegram
            caption="Berikut adalah hasil konversi tabel PDF dalam format JSON."
        )
        logger.info(f"File JSON PDF berhasil dikirim ke chat_id: {chat_id}")

    except Exception as e:
        logger.error(f"Gagal memproses PDF: {e}", exc_info=True)
        await context.bot.edit_message_text(
            text="❌ Terjadi kesalahan saat memproses PDF.",
            chat_id=chat_id,
            message_id=message_id
        )
    finally:
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)
            logger.info(f"Menghapus file sementara PDF: {temp_pdf_path}")
        if output_json_path and os.path.exists(output_json_path):
            # os.remove(output_json_path) # Biarkan terkomentar untuk debugging
            logger.info(f"File JSON output PDF tetap ada di: {output_json_path}")
        elif output_json_path:
            logger.warning(f"File JSON output PDF tidak ditemukan untuk dihapus: {output_json_path}")


async def handle_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    os.makedirs("temp_files", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    chat_id = update.effective_chat.id
    logger.info(f"Menerima PDF dari chat_id: {chat_id}")
    pdf_file = await update.message.document.get_file()
    
    # Ambil nama file asli dari dokumen yang diunggah
    original_filename = update.message.document.file_name
    # Dapatkan nama dasar tanpa ekstensi
    base_filename = os.path.splitext(original_filename)[0]

    temp_pdf_path = os.path.join("temp_files", f"{pdf_file.file_id}.pdf")
    await pdf_file.download_to_drive(temp_pdf_path)
    logger.info(f"PDF disimpan sementara di: {temp_pdf_path}")
    status_message = await context.bot.send_message(
        chat_id=chat_id,
        text="✅ File PDF diterima. Memulai ekstraksi tabel..."
    )
    context.application.create_task(
        process_pdf_and_send_json(context, chat_id, temp_pdf_path, status_message.message_id, base_filename)
    )

async def handle_docx(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    os.makedirs("temp_files", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    chat_id = update.effective_chat.id
    logger.info(f"Menerima DOCX dari chat_id: {chat_id}")
    docx_file = await update.message.document.get_file()
    
    # Ambil nama file asli dari dokumen yang diunggah
    original_filename = update.message.document.file_name
    # Dapatkan nama dasar tanpa ekstensi
    base_filename = os.path.splitext(original_filename)[0]

    temp_docx_path = os.path.join("temp_files", f"{docx_file.file_id}.docx")
    await docx_file.download_to_drive(temp_docx_path)
    logger.info(f"DOCX disimpan sementara di: {temp_docx_path}")
    status_message = await context.bot.send_message(
        chat_id=chat_id,
        text="✅ File DOCX diterima. Memulai ekstraksi tabel..."
    )
    output_json_path = None
    try:
        data = docx_to_json(temp_docx_path)
        data = fix_empty_key(data, new_key="Akun")
        if not data:
            await context.bot.edit_message_text(
                text="⚠️ Tidak ditemukan tabel pada DOCX.",
                chat_id=chat_id,
                message_id=status_message.message_id
            )
            logger.info("Tidak ada data tabel yang diekstrak dari DOCX.")
            return
        
        # Gunakan nama file asli sebagai nama file JSON
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        unique_id = uuid.uuid4().hex[:8] 
        output_json_path = os.path.join("output", f"{base_filename}_{timestamp}_{unique_id}.json")
        logger.info(f"Akan menyimpan JSON ke: {output_json_path}")

        with open(output_json_path, "w", encoding="utf-8") as f:
            json_to_write = json.dumps(data, ensure_ascii=False, indent=2)
            f.write(json_to_write)
        logger.info(f"JSON berhasil ditulis ke: {output_json_path}")

        await context.bot.edit_message_text(
            text="✅ JSON berhasil dibuat. Mengirim file ke Anda...",
            chat_id=chat_id,
            message_id=status_message.message_id
        )
        await context.bot.send_document(
            chat_id=chat_id,
            document=open(output_json_path, 'rb'),
            filename=os.path.basename(output_json_path), # Pastikan nama file benar untuk Telegram
            caption="Berikut adalah hasil konversi tabel DOCX dalam format JSON."
        )
        logger.info(f"File JSON DOCX berhasil dikirim ke chat_id: {chat_id}")

    except Exception as e:
        logger.error(f"Gagal memproses DOCX: {e}", exc_info=True)
        await context.bot.edit_message_text(
            text="❌ Terjadi kesalahan saat memproses DOCX.",
            chat_id=chat_id,
            message_id=status_message.message_id
        )
    finally:
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
            logger.info(f"Menghapus file sementara DOCX: {temp_docx_path}")
        if output_json_path and os.path.exists(output_json_path):
            # os.remove(output_json_path) # Biarkan terkomentar untuk debugging
            logger.info(f"File JSON output DOCX tetap ada di: {output_json_path}")
        elif output_json_path:
            logger.warning(f"File JSON output DOCX tidak ditemukan untuk dihapus: {output_json_path}")


async def handle_image(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    os.makedirs("temp_images", exist_ok=True)
    os.makedirs("output", exist_ok=True)

    chat_id = update.effective_chat.id
    logger.info(f"Menerima gambar dari chat_id: {chat_id}")
    photo_file = await update.message.photo[-1].get_file()
    
    # Untuk foto, Telegram tidak menyediakan nama file asli.
    # Kita akan menggunakan file_unique_id sebagai nama dasar.
    base_filename = photo_file.file_unique_id

    temp_image_path = os.path.join("temp_images", f"{photo_file.file_id}.jpg")
    await photo_file.download_to_drive(temp_image_path)
    logger.info(f"Gambar disimpan sementara di: {temp_image_path}")

    status_message = await context.bot.send_message(
        chat_id=chat_id,
        text="✅ Gambar diterima. Memulai analisis AI..."
    )

    context.application.create_task(
        process_image_and_send_json(context, chat_id, temp_image_path, status_message.message_id, base_filename)
    )


async def process_image_and_send_json(context: ContextTypes.DEFAULT_TYPE, chat_id: int, temp_image_path: str, message_id: int, original_base_filename: str):
    output_json_path = None

    try:
        await context.bot.edit_message_text(
            text="⏳ AI sedang memproses gambar untuk menghasilkan JSON...",
            chat_id=chat_id,
            message_id=message_id
        )

        json_result = ""
        logger.info(f"Memulai streaming JSON dari Gemini untuk gambar: {temp_image_path}")
        async for chunk in gemini_vision_extractor.stream_json_output(temp_image_path):
            json_result += chunk
        logger.info(f"Selesai streaming dari Gemini. Ukuran hasil: {len(json_result)} karakter.")

        print("DEBUG OUTPUT GEMINI (Raw):", json_result)

        if json_result.strip().startswith("```"):
            logger.info("Mendeteksi format markdown code block, membersihkan.")
            json_result = json_result.strip().lstrip("`json").lstrip("`").strip()
            if json_result.endswith("```"):
                json_result = json_result[:json_result.rfind("```")].strip()
            logger.info(f"Setelah membersihkan markdown. Ukuran hasil: {len(json_result)} karakter.")


        if not json_result.strip().startswith("["):
            await context.bot.edit_message_text(
                text="⚠️ Maaf, AI tidak dapat menghasilkan JSON dari gambar ini (hasil tidak dimulai dengan '[').",
                chat_id=chat_id,
                message_id=message_id
            )
            logger.warning(f"Hasil Gemini bukan JSON array: {json_result[:100]}...")
            return

        # --- Tambahan: Perbaiki key kosong ---
        data = []
        try:
            data = json.loads(json_result)
            logger.info("Berhasil parsing JSON dari hasil Gemini.")
            data = fix_empty_key(data, new_key="Akun")
            json_result_fixed = json.dumps(data, ensure_ascii=False, indent=2)
            logger.info("Berhasil memperbaiki key kosong dan memformat ulang JSON.")
        except json.JSONDecodeError as jde:
            logger.error(f"Gagal parsing JSON hasil Gemini: {jde}. Menggunakan hasil mentah.", exc_info=True)
            json_result_fixed = json_result.strip()
            await context.bot.edit_message_text(
                text=f"⚠️ Terjadi kesalahan saat memproses JSON dari AI. Coba lagi atau pastikan gambar tabel jelas. Error: {jde}",
                chat_id=chat_id,
                message_id=message_id
            )
            return
        except Exception as e:
            logger.error(f"Gagal memproses JSON hasil Gemini (selain JSONDecodeError): {e}", exc_info=True)
            json_result_fixed = json_result.strip()
            await context.bot.edit_message_text(
                text=f"⚠️ Terjadi kesalahan tidak terduga saat memproses JSON dari AI. Error: {e}",
                chat_id=chat_id,
                message_id=message_id
            )
            return
        # --------------------------------------

        # Gunakan nama file asli sebagai nama file JSON
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        unique_id = uuid.uuid4().hex[:8]
        output_json_path = os.path.join("output", f"{original_base_filename}_{timestamp}_{unique_id}.json")

        with open(output_json_path, "w", encoding="utf-8") as f:
            f.write(json_result_fixed)
        logger.info(f"JSON berhasil ditulis ke: {output_json_path}")

        await context.bot.edit_message_text(
            text="✅ JSON berhasil dibuat. Mengirim file ke Anda...",
            chat_id=chat_id,
            message_id=message_id
        )
        await context.bot.send_document(
            chat_id=chat_id,
            document=open(output_json_path, 'rb'),
            filename=os.path.basename(output_json_path), # Pastikan nama file benar untuk Telegram
            caption="Berikut adalah hasil konversi tabel dalam format JSON."
        )
        logger.info(f"File JSON gambar berhasil dikirim ke chat_id: {chat_id}")

    except Exception as e:
        logger.error(f"Gagal memproses gambar: {e}", exc_info=True)
        await context.bot.edit_message_text(
            text="❌ Terjadi kesalahan saat memproses gambar.",
            chat_id=chat_id,
            message_id=message_id
        )
    finally:
        if os.path.exists(temp_image_path):
            os.remove(temp_image_path)
            logger.info(f"Menghapus file sementara gambar: {temp_image_path}")
        if output_json_path and os.path.exists(output_json_path):
            # os.remove(output_json_path) # Biarkan terkomentar untuk debugging
            logger.info(f"File JSON output gambar tetap ada di: {output_json_path}")
        elif output_json_path:
            logger.warning(f"File JSON output gambar tidak ditemukan untuk dihapus: {output_json_path}")


def fix_empty_key(json_data, new_key="Akun"):
    if not json_data:
        return json_data
    old_key = ""
    if isinstance(json_data, list) and len(json_data) > 0 and old_key in json_data[0]:
        for obj in json_data:
            if isinstance(obj, dict) and old_key in obj:
                obj[new_key] = obj.pop(old_key)
    return json_data


def main() -> None:
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    application.add_handler(MessageHandler(filters.Document.PDF, handle_pdf))
    application.add_handler(MessageHandler(filters.Document.DOCX, handle_docx))
    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.PHOTO, handle_image))

    print("="*50)
    print("INFO: Bot berhasil dimulai dan siap menerima gambar.")
    print("="*50)
    logger.info("Bot starting polling...")
    application.run_polling()


if __name__ == "__main__":
    main()

